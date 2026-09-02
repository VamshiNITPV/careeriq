"""Text extraction from PDF and DOCX (ml.md section 2.1).

The output of this stage determines everything downstream: no text means no
sections, no skills, and an empty profile. So the failure modes are handled
explicitly rather than allowed to produce plausible-looking emptiness.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO

import pdfplumber
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import BadRequestError
from app.core.logging import get_logger
from app.services.file_validation import DocumentType

log = get_logger(__name__)

# Below this, the document is treated as containing no usable text. A scanned
# resume typically extracts a handful of stray characters rather than zero, so
# a strict `== 0` check would let it through and produce an empty profile the
# user cannot explain (ml.md section 2.1).
MIN_EXTRACTED_CHARACTERS = 100


class UnextractableDocumentError(BadRequestError):
    """The file is valid but contains no extractable text."""

    status_code = 422
    code = "UNEXTRACTABLE_DOCUMENT"
    message = (
        "No text could be read from this document. If it is a scan or an image, "
        "please upload a text-based PDF or a DOCX file."
    )


@dataclass(frozen=True, slots=True)
class ExtractedText:
    text: str
    page_count: int
    character_count: int
    extractor: str


def _normalize(text: str) -> str:
    """Clean extractor output without destroying structure.

    Line breaks are load-bearing here — section detection reads them — so this
    collapses noise while leaving paragraph boundaries intact.
    """
    # Ligatures. pdfplumber returns the actual Unicode ligature characters, and
    # leaving them turns "workflow" into a token no skill lookup will match.
    for ligature, replacement in (
        ("ﬀ", "ff"),
        ("ﬁ", "fi"),
        ("ﬂ", "fl"),
        ("ﬃ", "ffi"),
        ("ﬄ", "ffl"),
    ):
        text = text.replace(ligature, replacement)

    # Bullet glyphs vary by producer; normalise to a single marker so section
    # and list detection has one thing to look for.
    text = re.sub(r"[•●▪◦‣⁃·]", "•", text)

    # Non-breaking and exotic spaces read as ordinary spaces to a human and as
    # different characters to a tokenizer.
    text = text.replace(" ", " ").replace(" ", " ").replace("​", "")

    # Trailing whitespace per line, then collapse runs of blank lines to at most
    # one — PDFs frequently emit a blank line per visual row.
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _extract_pdf(content: bytes) -> ExtractedText:
    pages: list[str] = []
    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                # layout=True keeps column structure: without it a two-column
                # resume interleaves the columns line by line into nonsense.
                pages.append(page.extract_text(layout=True) or "")
            page_count = len(pdf.pages)
    except Exception as exc:
        # Encrypted, malformed, or a PDF feature pdfminer cannot handle.
        log.warning("pdf extraction failed", error=str(exc), error_type=type(exc).__name__)
        raise UnextractableDocumentError(
            "This PDF could not be read. It may be password protected or damaged."
        ) from exc

    return ExtractedText(
        text=_normalize("\n\n".join(pages)),
        page_count=page_count,
        character_count=0,  # set by the caller after normalisation
        extractor="pdfplumber",
    )


def _extract_docx(content: bytes) -> ExtractedText:
    try:
        document = DocxDocument(BytesIO(content))
    except (PackageNotFoundError, KeyError, ValueError) as exc:
        log.warning("docx extraction failed", error=str(exc))
        raise UnextractableDocumentError("This DOCX file could not be read.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]

    # Many resumes lay out skills and dates in tables. Ignoring tables silently
    # drops exactly the content most worth extracting.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ExtractedText(
        text=_normalize("\n".join(parts)),
        page_count=1,  # DOCX has no fixed pagination
        character_count=0,
        extractor="python-docx",
    )


def extract_text(*, content: bytes, document_type: DocumentType) -> ExtractedText:
    """Extract plain text, or raise UnextractableDocumentError."""
    result = _extract_pdf(content) if document_type is DocumentType.PDF else _extract_docx(content)

    character_count = len(result.text)
    if character_count < MIN_EXTRACTED_CHARACTERS:
        log.info(
            "document produced too little text",
            characters=character_count,
            minimum=MIN_EXTRACTED_CHARACTERS,
            extractor=result.extractor,
        )
        raise UnextractableDocumentError()

    return ExtractedText(
        text=result.text,
        page_count=result.page_count,
        character_count=character_count,
        extractor=result.extractor,
    )
