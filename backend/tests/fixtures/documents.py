"""Builders for real PDF and DOCX test documents.

Generated rather than checked in as binaries: a committed fixture is opaque in
review, and hand-crafted PDF bytes would test our own byte-writing rather than
the extraction path a genuine file exercises.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from fpdf import FPDF

# A deliberately ordinary resume. The skills are spread across three sections so
# tests can assert that section-aware confidence actually differs, and it
# includes aliases ("Postgres", "K8s", "JS") to prove alias resolution works
# rather than only exact names matching.
SAMPLE_RESUME = """\
PRIYA SHARMA
priya.sharma@example.com | +91 98765 43210 | Bengaluru, India
github.com/priyasharma | linkedin.com/in/priyasharma

PROFESSIONAL SUMMARY
Backend engineer with 3 years of experience building REST APIs and
distributed services. Comfortable owning a feature from schema design
through deployment.

TECHNICAL SKILLS
Languages: Python, JavaScript, SQL, Go
Frameworks: FastAPI, Django, React
Databases: Postgres, Redis, MongoDB
Cloud & Tools: Docker, K8s, AWS, Git, CI/CD

WORK EXPERIENCE
Backend Engineer, Zenith Systems
June 2023 - Present
- Built and maintained REST APIs in Python serving 2M requests per day
- Migrated a monolith to microservices, reducing deploy time
- Introduced pytest coverage gates into the CI pipeline

Software Engineering Intern, Nova Labs
January 2023 - May 2023
- Implemented a caching layer with Redis
- Wrote integration tests for the payments service

EDUCATION
B.Tech in Computer Science
National Institute of Technology
2019 - 2023

PROJECTS
Resume Matcher
Semantic search over job descriptions using embeddings and pgvector.
Built with FastAPI and React.

CERTIFICATIONS
AWS Certified Cloud Practitioner, 2024
"""

# No recognisable headings at all — exercises the fallback path where section
# detection finds nothing and extraction has to scan the whole document.
UNSTRUCTURED_RESUME = """\
Priya Sharma is a backend engineer based in Bengaluru who has spent the last
three years writing Python and TypeScript. She works with PostgreSQL and Redis
daily, deploys with Docker, and has recently been learning Kubernetes. Before
that she interned at Nova Labs where she built REST APIs and wrote tests
using pytest. She holds a B.Tech in Computer Science and enjoys mentoring
junior engineers on system design and code review practices.
"""


def build_pdf(text: str = SAMPLE_RESUME) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    for line in text.splitlines():
        # Latin-1 is the built-in font's encoding; the sample text is ASCII, so
        # this only guards against an accidental smart quote breaking the build.
        safe = line.encode("latin-1", "replace").decode("latin-1")
        pdf.cell(0, 5, safe, new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


def build_docx(text: str = SAMPLE_RESUME) -> bytes:
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_docx_with_table() -> bytes:
    """A resume that puts its skills in a table.

    Common in real resumes and easy to miss: python-docx's `paragraphs` does not
    include table cells, so an extractor that reads only paragraphs silently
    drops exactly the section most worth extracting.
    """
    document = Document()
    document.add_paragraph("PRIYA SHARMA")
    document.add_paragraph("priya.sharma@example.com")
    document.add_paragraph("TECHNICAL SKILLS")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Languages"
    table.cell(0, 1).text = "Python, TypeScript, SQL"
    table.cell(1, 0).text = "Databases"
    table.cell(1, 1).text = "PostgreSQL, Redis"
    table.cell(2, 0).text = "Cloud"
    table.cell(2, 1).text = "AWS, Docker, Kubernetes"

    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("Backend Engineer at Zenith Systems, 2023 to present.")
    document.add_paragraph("Built REST APIs and maintained CI/CD pipelines for the team.")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("B.Tech in Computer Science, NIT, 2019 to 2023.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_image_only_pdf() -> bytes:
    """A valid PDF containing no extractable text.

    Stands in for a scanned resume — the case that must be rejected with a clear
    message rather than silently producing an empty profile.
    """
    pdf = FPDF()
    pdf.add_page()
    return bytes(pdf.output())
